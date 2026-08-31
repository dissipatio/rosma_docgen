"""
rosma_docgen — Stage 1 (infrastructure only)

Purpose of this stage: prove that LibreOffice + Cyrillic fonts + DOCX->PDF
conversion actually work on Railway, BEFORE any real template exists.

What works now:
  GET  /health        - liveness
  GET  /fonts         - which fonts LibreOffice can actually see
  POST /smoke-test    - builds a Cyrillic DOCX, converts it, returns the PDF
  POST /render-test   - upload your own .docx template + JSON context, get PDF back
  POST /generate      - stub; wired up in Stage 2 once field mapping is done

Stage 2 adds: Airtable fetch, template lookup, context builder, upload-back.
"""

import json
import logging
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, JSONResponse

from pdf_convert import docx_to_pdf

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
log = logging.getLogger("docgen")

app = FastAPI(title="ROSMA docgen")


# --------------------------------------------------------------------------
# Diagnostics
# --------------------------------------------------------------------------

@app.get("/health")
def health():
    soffice = shutil.which("soffice")
    return {
        "status": "ok",
        "soffice": soffice,
        "soffice_found": bool(soffice),
    }


@app.get("/fonts")
def fonts():
    """
    Lists the font families fontconfig exposes to LibreOffice.

    If 'Liberation Serif' is missing here, Cyrillic text and Times New Roman
    substitution will both be wrong in the output PDF.
    """
    try:
        out = subprocess.run(
            ["fc-list", ":", "family"],
            capture_output=True, text=True, timeout=30,
        ).stdout
    except Exception as e:
        raise HTTPException(500, f"fc-list failed: {e}")

    families = sorted({
        part.strip()
        for line in out.splitlines()
        for part in line.split(",")
        if part.strip()
    })

    expected = ["Liberation Serif", "Liberation Sans", "DejaVu Sans"]
    return {
        "total": len(families),
        "expected_present": {name: name in families for name in expected},
        "families": families,
    }


# --------------------------------------------------------------------------
# Smoke test — no template needed
# --------------------------------------------------------------------------

@app.post("/smoke-test")
def smoke_test():
    """
    Builds a small DOCX containing Cyrillic text in Times New Roman,
    converts it to PDF, and returns the PDF.

    Open the result and check: Cyrillic renders as letters (not boxes),
    and the serif font looks like Times, not a fallback.
    """
    from docx import Document
    from docx.shared import Pt
    from num2words import num2words

    tmp = Path(tempfile.mkdtemp())
    docx_path = tmp / "smoke.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading("Проверка кириллицы", level=1)
    doc.add_paragraph("Счёт на оплату № 123 от 12 августа 2026 г.")
    doc.add_paragraph("Поставщик: ООО «РОСМА» — штампы, бандажи, реставрация.")
    doc.add_paragraph("ЁёЙйЩщЪъЫыЬьЭэЮюЯя — проверка редких букв.")
    doc.add_paragraph(
        "Сумма прописью: " + num2words(1234.56, to="currency", lang="ru")
    )

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Наименование", "Кол-во", "Сумма"]):
        table.rows[0].cells[i].text = h
    for i, v in enumerate(["Штамп ГРАФ-200", "2", "45 000,00"]):
        table.rows[1].cells[i].text = v

    doc.save(docx_path)

    try:
        pdf_path = docx_to_pdf(docx_path, tmp)
    except RuntimeError as e:
        raise HTTPException(500, str(e))

    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        filename="smoke_test.pdf",
    )


@app.post("/render-test")
async def render_test(
    template: UploadFile = File(...),
    context: str = Form("{}"),
):
    """
    Upload a real .docx template plus a JSON context, get the rendered PDF back.

    This is how we test each Bitrix template as it gets converted, without
    touching Airtable at all.

    curl -X POST $URL/render-test \
      -F "template=@schet.docx" \
      -F 'context={"client_name":"ООО Тест","total":45000}' \
      -o out.pdf
    """
    from docxtpl import DocxTemplate

    try:
        ctx = json.loads(context)
    except json.JSONDecodeError as e:
        raise HTTPException(400, f"context is not valid JSON: {e}")

    if not template.filename.lower().endswith(".docx"):
        raise HTTPException(400, "template must be a .docx file")

    tmp = Path(tempfile.mkdtemp())
    tpl_path = tmp / "template.docx"
    tpl_path.write_bytes(await template.read())

    out_docx = tmp / "rendered.docx"
    try:
        tpl = DocxTemplate(tpl_path)
        tpl.render(ctx)
        tpl.save(out_docx)
    except Exception as e:
        raise HTTPException(400, f"docxtpl render failed: {e}")

    try:
        pdf_path = docx_to_pdf(out_docx, tmp)
    except RuntimeError as e:
        raise HTTPException(500, str(e))

    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        filename="rendered.pdf",
    )


# --------------------------------------------------------------------------
# Stage 2 — real Airtable-driven generation
# --------------------------------------------------------------------------

@app.post("/generate")
@app.post("/generate-document")  # alias — matches the Airtable Automation's webhook URL
async def generate(payload: dict, background_tasks: BackgroundTasks):
    """
    Airtable-triggered generation, fired by a generation checkbox on
    Inquiries or (Договор поставки onward) Clients. Returns 202 immediately
    because LibreOffice conversion + the Yandex Disk upload take longer
    than Airtable's webhook timeout — generate_document_for_record()
    writes its own result (link, status, error) back to the record
    regardless of this response.

    payload:
      record_id      required. The triggering record.
      table_id       optional. Defaults to Inquiries for backward
                     compatibility with the original automation, which
                     only ever sends record_id. Any other scope table
                     (e.g. Clients) must send this explicitly.
      template_name  optional. Only needed for scope tables that don't
                     have their own "Шаблон для генерации" link field —
                     Inquiries resolves its template from that link as
                     before; Clients doesn't have one yet (single-purpose
                     trigger, Договор поставки only), so its automation
                     sends template_name directly instead.
    """
    from generate_and_deliver import generate_document_for_record

    record_id = payload.get("record_id")
    if not record_id:
        raise HTTPException(400, "record_id is required")

    log.info("Generate requested for %s (table=%s, template=%s)",
              record_id, payload.get("table_id"), payload.get("template_name"))
    background_tasks.add_task(
        generate_document_for_record,
        record_id,
        table_id=payload.get("table_id"),
        template_name=payload.get("template_name"),
    )

    return JSONResponse(
        status_code=202,
        content={"accepted": True, "record_id": record_id},
    )
